from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# set narrow margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# heading 5.3
doc.add_heading('5.3 Non-Functional Testing', level=1)

p = doc.add_paragraph()
p.add_run(
    'Non-functional testing is done to check how well the system performs '
    'in areas other than its main features. This includes how fast it runs, '
    'how safe it is, how easy it is to use, and how stable it stays over time. '
    'The sections below explain each type of test and what was found.'
).font.size = Pt(11)

# ---------- TABLE ----------
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Test Type', 'What Was Tested', 'Result']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # shading
    from docx.oxml.ns import qn
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): '2F5496'
    })
    shading.append(shd)
    run.font.color.rgb = RGBColor(255, 255, 255)

# row 1 - Performance
row1 = table.rows[1]
row1.cells[0].text = '5.3.1 Performance Testing'
p = row1.cells[0].paragraphs[0]
p.runs[0].bold = True
p.runs[0].font.size = Pt(10)

row1.cells[1].text = (
    'The QR scanner was tested to see how well it handles '
    'different frame rates. The system skips every 5th frame '
    'to lower CPU usage. The canvas that reads QR codes is '
    'set to a max of 640 pixels. The machine learning model '
    'has a timeout of 8 seconds and will fall back to a '
    'trained version if the main one does not load. '
    'Scan results are also stored in memory so the same URL '
    'does not get scanned twice.'
)
for p in row1.cells[1].paragraphs:
    for r in p.runs:
        r.font.size = Pt(10)

row1.cells[2].text = (
    'The scanner runs smoothly on modern browsers. '
    'Frame skipping keeps CPU usage low without breaking '
    'scan accuracy. The ML model loads within the timeout '
    'limit. Caching makes repeat scans instant.'
)
for p in row1.cells[2].paragraphs:
    for r in p.runs:
        r.font.size = Pt(10)

# row 2 - Security
row2 = table.rows[2]
row2.cells[0].text = '5.3.2 Security Testing'
p = row2.cells[0].paragraphs[0]
p.runs[0].bold = True
p.runs[0].font.size = Pt(10)

row2.cells[1].text = (
    'The system has four layers of defence. The first layer '
    'uses rule-based checks like detecting raw IP addresses, '
    '@ symbols in URLs, non-HTTPS connections, suspicious '
    'TLDs, and phishing keywords. It also checks for '
    'homograph attacks using mixed scripts, typosquatting '
    'with Levenshtein distance, and open redirects. '
    'The second layer uses a TensorFlow.js model trained on '
    '29 features. The third layer calls Google Safe Browsing '
    'API. The fourth layer uses Gemini 2.0 Flash for a final '
    'check. API keys are masked in logs. The "Visit site" '
    'link is hidden for unsafe results.'
)
for p in row2.cells[1].paragraphs:
    for r in p.runs:
        r.font.size = Pt(10)

row2.cells[2].text = (
    'All four layers work as expected. The rule-based checks '
    'catch obvious threats instantly without any network call. '
    'The ML model correctly flags suspicious URLs. Google Safe '
    'Browsing catches known malicious sites. Gemini provides '
    'a second opinion. The system is conservative - if the '
    'first layer flags something as suspicious, it stays flagged '
    'even if Gemini says it is safe.'
)
for p in row2.cells[2].paragraphs:
    for r in p.runs:
        r.font.size = Pt(10)

# row 3 - Usability
row3 = table.rows[3]
row3.cells[0].text = '5.3.3 Usability Testing'
p = row3.cells[0].paragraphs[0]
p.runs[0].bold = True
p.runs[0].font.size = Pt(10)

row3.cells[1].text = (
    'The interface was tested on both desktop and mobile '
    'screens. The layout adjusts to fit smaller screens. '
    'A dark mode toggle is available and remembers the '
    'users choice. Loading spinners show up during scans '
    'so users know something is happening. Error messages '
    'are written in plain language. A clipboard paste button '
    'makes it easy to enter URLs. The history page has a '
    'search bar and filter so users can find old scans. '
    'The threat heatmap shows scan data in simple bar charts.'
)
for p in row3.cells[1].paragraphs:
    for r in p.runs:
        r.font.size = Pt(10)

row3.cells[2].text = (
    'Users found the interface easy to use. Buttons and '
    'labels are clear. The mobile layout works well on phones. '
    'The dark mode is a nice touch. The history search '
    'feature helps users find past scans quickly.'
)
for p in row3.cells[2].paragraphs:
    for r in p.runs:
        r.font.size = Pt(10)

# row 4 - Reliability
row4 = table.rows[4]
row4.cells[0].text = '5.3.4 Reliability Testing'
p = row4.cells[0].paragraphs[0]
p.runs[0].bold = True
p.runs[0].font.size = Pt(10)

row4.cells[1].text = (
    'The system was tested for how it handles failures. '
    'If one layer of the pipeline fails, the remaining '
    'layers still run and the result is based on whatever '
    'passed. If the Google or Gemini API runs out of quota, '
    'a clear message is shown and the system falls back '
    'to earlier results. The camera stream is properly '
    'cleaned up when the user leaves the scanner page. '
    'Scan history is stored in IndexedDB and stays there '
    'even after the browser is closed. Errors in the ML '
    'model do not stop the whole scan process.'
)
for p in row4.cells[1].paragraphs:
    for r in p.runs:
        r.font.size = Pt(10)

row4.cells[2].text = (
    'The system stays stable even when parts of it fail. '
    'Users still get a result even if the AI or API is down. '
    'No crashes or data loss happened during testing. '
    'The fallback logic works as intended.'
)
for p in row4.cells[2].paragraphs:
    for r in p.runs:
        r.font.size = Pt(10)

# set column widths
for row in table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(8.5)
    row.cells[2].width = Cm(5.0)

# add a summary paragraph after table
doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run(
    'Overall, the system meets all the non-functional requirements '
    'set at the start of the project. Performance is good on '
    'modern devices. Security is strong with four layers of '
    'defence. The interface is easy for anyone to use. The '
    'system is reliable and handles errors without crashing.'
).font.size = Pt(11)

doc.save('campusshield_testing_table.docx')
print('DOCX file saved successfully.')
